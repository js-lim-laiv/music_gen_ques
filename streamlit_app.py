import streamlit as st
from io import BytesIO
from docx import Document
import random
from transformers import pipeline

# LLM λ¨λΈ μ΄κΈ°ν™” (ν•κµ­μ–΄ νΉν™” λ¨λΈ μ‚¬μ©)
try:
    text_generator = pipeline("text-generation", model="beomi/KoAlpaca-Polyglot-12.8B")
    gpt_ready = True
except:
    gpt_ready = False

# Streamlit UI
st.set_page_config(page_title="AI μμ•… λ¬Έν•­ μƒμ„±κΈ°", layout="wide")
st.title("πΌ AI κΈ°λ° μμ•… λ¬Έν•­ μƒμ„±κΈ° (PoC)")
st.markdown("---")

col1, col2 = st.columns([1, 3])

with col1:
    st.subheader("1οΈβƒ£ μƒμ„±ν•  λ¬Έν•­ μ ν•")
    question_type = st.selectbox("λ¬Έν•­ μ ν•", [
        "μ ν• 1: μμ•…μ‚¬ (LLM κΈ°λ° μ‹¤μƒμ„±)",
        "μ ν• 2: λ¦¬λ“¬/ν™”μ„± (νμΌ κΈ°λ° λ¨μ μƒμ„±)",
        "μ ν• 3: μ•…λ³΄ ν‰κ°€ (νμΌλ… κΈ°λ° λ¬Έν•­ μƒμ„±)",
        "μ ν• 4: μΆ…ν•© ν‰κ°€"
    ])

    st.subheader("2οΈβƒ£ μ •λ‹µ μ ν•")
    answer_type = st.selectbox("μ •λ‹µ μ ν•", ["O/X", "κ°κ΄€μ‹ (ν…μ¤νΈ)", "μ•…λ³΄ν• λ³΄κΈ°", "μ„μ ν•"])

    audio_file = st.file_uploader("πµ μ¤λ””μ¤ μ—…λ΅λ“ (wav/mp3)", type=["wav", "mp3"])
    score_file = st.file_uploader("πΌ μ•…λ³΄ μ—…λ΅λ“ (musicxml)", type=["xml", "musicxml"])
    generate = st.button("β¨ λ¬Έν•­ μƒμ„±ν•κΈ°")

with col2:
    st.subheader("π§Ύ μƒμ„± κ²°κ³Ό")
    st.caption("μ•„λλ” μ‹¤μ  λ¨λΈ κΈ°λ°μΌλ΅ μƒμ„±λ λ¬Έν•­μ…λ‹λ‹¤.")

    if audio_file:
        st.audio(audio_file, format='audio/wav' if audio_file.name.endswith(".wav") else 'audio/mp3')

    def generate_llm_question(prompt):
        if not gpt_ready:
            return "β οΈ λ¨λΈμ΄ μ¤€λΉ„λμ§€ μ•μ•„ μμ‹ λ¬Έν•­μ„ μ‚¬μ©ν•©λ‹λ‹¤."
        try:
            result = text_generator(prompt, max_length=300, do_sample=True, temperature=0.8)[0]["generated_text"]
            return result
        except Exception as e:
            return f"β οΈ LLM μƒμ„± μ¤λ¥: {str(e)}"

    def rhythm_mock():
        return f"Q. μ΄ μμ›μ λ¦¬λ“¬ μ ν•μ€ λ¬΄μ—‡μΈκ°€μ”?\nμ •λ‹µ: {random.choice(['μ™μΈ ', 'λ³΄μ‚¬λ…Έλ°”', 'ν‘ν¬', 'λ””μ¤μ½”'])}"

    def score_mock(filename):
        name = filename.name.replace(".musicxml", "").replace(".xml", "")
        return f"Q. μ•…λ³΄ '{name}'μ€ μ–΄λ–¤ μ΅°μ„±μ„ κΈ°λ°μΌλ΅ λ§λ“¤μ–΄μ΅μ„κΉμ”?\nμ •λ‹µ: {random.choice(['λ‹¤μ¥μ΅°', 'κ°€λ‹¨μ΅°', 'λ°”μ¥μ΅°'])}"

    def combine_all(text, audio, score):
        return (
            f"Q. λ‹¤μ κ³΅μ— λ€ν• λ¶„μ„ κ²°κ³Όλ΅ μ–΄λ–¤ μ‹λ€μ μμ•…μΌκΉμ”?\n"
            f"- πµ λ¦¬λ“¬ λ¶„μ„ κ²°κ³Ό: {audio}\n"
            f"- πΌ μ΅°μ„± λ¶„μ„ κ²°κ³Ό: {score}\n"
            f"- π“– λ¬Έν— κΈ°λ° μ •λ³΄: {text}\n"
            f"μ •λ‹µ: κ³ μ „μ£Όμ"
        )

    if generate:
        if "μμ•…μ‚¬" in question_type:
            # μ •λ‹µ μ ν• κΈ°λ° ν”„λ΅¬ν”„νΈ μ„¤μ •
            if answer_type == "O/X":
                prompt = "κ³ μ „ μμ•…μ‚¬μ™€ κ΄€λ ¨λ O/X λ¬Έμ λ¥Ό ν•κµ­μ–΄λ΅ λ§λ“¤μ–΄μ¤. μ§λ¬Έκ³Ό μ •λ‹µ(O λλ” X)μ„ λ…ν™•ν•κ² ν¬ν•¨ν•΄μ¤."
            elif "κ°κ΄€μ‹" in answer_type:
                prompt = ("κ³ μ „ μμ•…μ‚¬ κ΄€λ ¨ κ°κ΄€μ‹ λ¬Έμ λ¥Ό ν•κµ­μ–΄λ΅ μƒμ„±ν•΄μ¤. "
                          "λ³΄κΈ° A) B) C) D) ν•μ‹κ³Ό μ •λ‹µ ν‘μ‹λ„ ν¬ν•¨ν•΄μ¤.")
            elif "μ„μ ν•" in answer_type:
                prompt = "κ³ μ „ μμ•…μ‚¬μ— λ€ν• μ„μ ν• λ¬Έν•­μ„ μƒμ„±ν•΄μ¤. μ§λ¬Έμ€ λ…ν™•ν•κ², μ •λ‹µ μμ‹λ” ν• λ¬Έμ¥μΌλ΅."
            else:
                prompt = "κ³ μ „ μμ•…μ‚¬ κ΄€λ ¨ λ¬Έν•­μ„ μƒμ„±ν•΄μ¤."

            result = generate_llm_question(prompt)

        elif "λ¦¬λ“¬" in question_type and audio_file:
            result = rhythm_mock()

        elif "μ•…λ³΄" in question_type and score_file:
            result = score_mock(score_file)

        elif "μΆ…ν•©" in question_type:
            text = generate_llm_question("μμ•…μ‚¬ κ΄€λ ¨ μ„¤λ…μ„ μƒμ„±ν•΄μ¤. ν• λ¬Έλ‹¨ μ •λ„λ΅.")
            audio = rhythm_mock() if audio_file else "μ¤λ””μ¤ μ—†μ"
            score = score_mock(score_file) if score_file else "μ•…λ³΄ μ—†μ"
            result = combine_all(text, audio, score)

        else:
            result = "β οΈ μ…λ ¥μ΄ λ¶€μ΅±ν•κ±°λ‚ μ¬λ°”λ¥΄μ§€ μ•μµλ‹λ‹¤."

        st.text_area("π“ μƒμ„±λ λ¬Έν•­", result, height=300)

        doc = Document()
        doc.add_heading("AI λ¬Έν•­ κ²°κ³Ό", 0)
        doc.add_paragraph(result)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button("β¬‡ Word λ‹¤μ΄λ΅λ“", data=buffer, file_name="music_question.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.info("λ¬Έν•­ μ ν•μ„ μ„ νƒν•κ³  [λ¬Έν•­ μƒμ„±ν•κΈ°]λ¥Ό λλ¬λ³΄μ„Έμ”.")
